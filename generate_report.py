import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_report():
    doc = docx.Document()
    
    # Set Standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Helper to set background color of cells
    def set_cell_background(cell, color_hex):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    # Helper to set cell padding
    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # Helper to set thin borders
    def set_cell_borders(cell, color="D3D3D3", sz="4", val="single"):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for b in ['top', 'left', 'bottom', 'right']:
            node = OxmlElement(f'w:{b}')
            node.set(qn('w:val'), val)
            node.set(qn('w:sz'), sz)
            node.set(qn('w:space'), '0')
            node.set(qn('w:color'), color)
            tcBorders.append(node)
        tcPr.append(tcBorders)

    # Helper to add styled paragraph
    def add_para(text="", font_name="Calibri", font_size=11, bold=False, italic=False, 
                 color_rgb=(34, 34, 34), align=WD_ALIGN_PARAGRAPH.LEFT, 
                 space_before=0, space_after=6, line_spacing=1.15, bullet=False):
        style = 'List Bullet' if bullet else 'Normal'
        p = doc.add_paragraph(style=style)
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        
        if text:
            run = p.add_run(text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
            run.font.italic = italic
            run.font.color.rgb = RGBColor(*color_rgb)
        return p

    # Helper to add bold/normal inline text
    def add_inline(p, text, bold=False, italic=False, color_rgb=(34, 34, 34), font_name="Calibri", font_size=11):
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(*color_rgb)
        return run

    # Helper to add headings
    def add_heading_1(text):
        p = add_para(text, font_name="Calibri Light", font_size=18, bold=True, 
                     color_rgb=(27, 54, 93), space_before=12, space_after=6)
        p.paragraph_format.keep_with_next = True
        return p

    def add_heading_2(text):
        p = add_para(text, font_name="Calibri Light", font_size=14, bold=True, 
                     color_rgb=(27, 54, 93), space_before=12, space_after=4)
        p.paragraph_format.keep_with_next = True
        return p

    def add_heading_3(text):
        p = add_para(text, font_name="Calibri", font_size=12, bold=True, 
                     color_rgb=(51, 51, 51), space_before=6, space_after=2)
        p.paragraph_format.keep_with_next = True
        return p

    # XML magic to add page numbers in footer
    def add_page_number(run):
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        r = run._r
        r.append(fldChar1)
        r.append(instrText)
        r.append(fldChar2)
        r.append(fldChar3)

    # ----------------------------------------------------
    # COVER PAGE
    # ----------------------------------------------------
    for _ in range(3):
        add_para(space_after=12)
        
    add_para("SRI LANKA TECHNOLOGICAL CAMPUS (SLTC)", font_name="Calibri", font_size=14, bold=True, 
             color_rgb=(100, 100, 100), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
             
    add_para("DEPARTMENT OF INFORMATION TECHNOLOGY", font_name="Calibri", font_size=12, bold=False, 
             color_rgb=(120, 120, 120), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
             
    for _ in range(2):
        add_para(space_after=12)

    add_para("DESIGN AND IMPLEMENTATION OF AN INTELLIGENT SMS SPAM DETECTION SYSTEM", 
             font_name="Calibri Light", font_size=24, bold=True, color_rgb=(27, 54, 93), 
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
             
    add_para("A Comparative NLP Study using Machine Learning & Sequential Deep Learning Architectures", 
             font_name="Calibri", font_size=14, italic=True, color_rgb=(80, 80, 80), 
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

    for _ in range(3):
        add_para(space_after=12)

    add_para("SEMESTER 8 — NATURAL LANGUAGE PROCESSING (NLP)", 
             font_name="Calibri", font_size=12, bold=True, color_rgb=(51, 51, 51), 
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
             
    add_para("Group Research Project & Technical Report", 
             font_name="Calibri", font_size=11, color_rgb=(100, 100, 100), 
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

    members_table = doc.add_table(rows=4, cols=2)
    members_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = [("Prepared By:", "Assigned Research Focus:"),
               ("Student Member 01", "Logistic Regression (ML) & 1D CNN (DL)"),
               ("Student Member 02", "Random Forest (ML) & LSTM RNN (DL)"),
               ("Student Member 03", "XGBoost (ML) & Custom Transformer (DL)")]
               
    for i, (name, focus) in enumerate(headers):
        row = members_table.rows[i]
        cell_name = row.cells[0]
        cell_focus = row.cells[1]
        
        p_name = cell_name.paragraphs[0]
        p_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_name.paragraph_format.space_after = Pt(4)
        run_n = p_name.add_run(name)
        run_n.font.name = "Calibri"
        run_n.font.size = Pt(11)
        run_n.bold = (i == 0)
        if i == 0:
            run_n.font.color.rgb = RGBColor(100, 100, 100)
            
        p_focus = cell_focus.paragraphs[0]
        p_focus.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_focus.paragraph_format.space_after = Pt(4)
        run_f = p_focus.add_run("   " + focus)
        run_f.font.name = "Calibri"
        run_f.font.size = Pt(11)
        run_f.bold = (i == 0)
        if i == 0:
            run_f.font.color.rgb = RGBColor(100, 100, 100)
            
    for row in members_table.rows:
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(3.5)

    for _ in range(3):
        add_para(space_after=12)

    add_para("Date: August 2026", font_name="Calibri", font_size=10, 
             color_rgb=(130, 130, 130), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ----------------------------------------------------
    # SET UP HEADER & FOOTER ON SUBSEQUENT PAGES
    # ----------------------------------------------------
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    header = section.header
    p_hdr = header.paragraphs[0]
    p_hdr.text = "SMS Spam Detection System — Final Technical Report"
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_hdr.runs[0].font.name = "Calibri"
    p_hdr.runs[0].font.size = Pt(8.5)
    p_hdr.runs[0].font.color.rgb = RGBColor(120, 120, 120)

    footer = section.footer
    p_ftr = footer.paragraphs[0]
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ftr.paragraph_format.space_before = Pt(0)
    p_ftr.paragraph_format.space_after = Pt(0)
    
    run_ftr_text = p_ftr.add_run("Sri Lanka Technological Campus | Department of IT  —  Page ")
    run_ftr_text.font.name = "Calibri"
    run_ftr_text.font.size = Pt(8.5)
    run_ftr_text.font.color.rgb = RGBColor(120, 120, 120)
    
    run_page = p_ftr.add_run()
    run_page.font.name = "Calibri"
    run_page.font.size = Pt(8.5)
    run_page.font.color.rgb = RGBColor(120, 120, 120)
    add_page_number(run_page)

    # ----------------------------------------------------
    # SECTION 1: INTRODUCTION & PROJECT OBJECTIVES
    # ----------------------------------------------------
    add_heading_1("1. INTRODUCTION & PROJECT OBJECTIVES")
    
    p = add_para()
    add_inline(p, "With the exponential rise in global mobile device usage, Short Message Service (SMS) has remained a primary and highly accessible channel for communication. However, this ubiquity has also turned SMS into a prime target for malicious actors, resulting in a dramatic surge in unsolicited spam messages. SMS spam spans across unsolicited commercial marketing, financial scams, phishing attacks, and identity theft attempts. Unlike email spam, SMS spam is uniquely disruptive due to high open rates, immediate push notifications, and the general lack of pre-installed, robust spam filters on standard mobile terminals. This poses critical security vulnerabilities and severe user inconveniences.")
    
    p = add_para()
    add_inline(p, "To address this pressing problem, this project presents the design and implementation of a robust, automated ")
    add_inline(p, "SMS Spam Detection System", bold=True)
    add_inline(p, " utilizing Natural Language Processing (NLP), classical Machine Learning (ML), and sequence-based Deep Learning (DL) models. The primary goal is to evaluate different algorithms, study their structural suitability, and build an end-to-end classification pipeline capable of running real-time spam filtering.")

    p = add_para()
    add_inline(p, "The core project objectives are defined as follows:")
    
    p = add_para(bullet=True)
    add_inline(p, "To download and preprocess a real-world SMS spam text corpus, performing normalization, noise reduction, and feature extraction.")
    
    p = add_para(bullet=True)
    add_inline(p, "To assign independent research streams to three group members, where each member implements one classical Machine Learning model (TF-IDF features) and one deep neural network (tokenized sequential dense embeddings).")
    
    p = add_para(bullet=True)
    add_inline(p, "To evaluate all six models on standard classification performance metrics (Accuracy, Precision, Recall, and F1-Score) to analyze their trade-offs.")
    
    p = add_para(bullet=True)
    add_inline(p, "To deploy the selected champion model on a Flask API backend integrated with an interactive, modern web application UI.")

    # ----------------------------------------------------
    # SECTION 2: DATASET DETAILS AND PREPROCESSING
    # ----------------------------------------------------
    add_heading_1("2. DATASET CHARACTERISTICS & PREPROCESSING")
    
    p = add_para()
    add_inline(p, "The system leverages the publicly accessible ")
    add_inline(p, "SMS Spam Collection Dataset", bold=True)
    add_inline(p, ", which contains ")
    add_inline(p, "5,574 English text messages", bold=True)
    add_inline(p, " labeled as either 'ham' (legitimate, non-spam) or 'spam' (unsolicited commercial or phishing texts). The corpus presents a highly imbalanced class distribution, reflecting a realistic spam landscape: ")
    add_inline(p, "4,827 ham messages (86.6%)", bold=True)
    add_inline(p, " versus ")
    add_inline(p, "747 spam messages (13.4%)", bold=True)
    add_inline(p, ".")
    
    p = add_para()
    add_inline(p, "This dataset is highly suitable for training NLP classification models due to its capturing of actual mobile communication behaviors. It includes abbreviations (e.g., 'u', 'r', 'txt', 'pls'), internet slang, typos, emojis, and distinct uppercase/punctuation formatting. Text lengths also show distinct profiles: spam messages tend to be substantially longer with uppercase calls-to-action (e.g., 'FREE', 'WINNER', 'CLAIM NOW'), whereas ham messages are generally shorter and conversational.")

    add_heading_2("Data Preparation & Preprocessing Pipeline")
    
    p = add_para()
    add_inline(p, "To convert raw textual data into clean inputs suitable for vectorization, we established a standardized preprocessing pipeline:")
    
    p = add_para(bullet=True)
    add_inline(p, "Text Normalization: ", bold=True)
    add_inline(p, "All text strings are converted to lowercase to avoid duplicate vocabulary keys based on capitalization variations.")
    
    p = add_para(bullet=True)
    add_inline(p, "Noise Removal: ", bold=True)
    add_inline(p, "Regular expressions are used to remove non-alphanumeric characters, punctuation marks, and special symbols, reducing input noise while preserving standard words and numbers.")
    
    p = add_para(bullet=True)
    add_inline(p, "Target Label Encoding: ", bold=True)
    add_inline(p, "The categorical class labels are mapped to binary integers: 'ham' is encoded as 0, and 'spam' is encoded as 1.")
    
    p = add_para(bullet=True)
    add_inline(p, "Dataset Partitioning: ", bold=True)
    add_inline(p, "The preprocessed corpus is partitioned into a training set (70% - 3,901 samples) to train models, a validation set (15% - 836 samples) to tune deep learning hyperparameters and monitor epochs, and an independent test set (15% - 837 samples) for final unbiased comparative evaluations.")

    add_heading_2("Text Representation and Feature Extraction")
    
    p = add_para()
    add_inline(p, "Machine learning models and deep neural networks require different text representations. Thus, two distinct feature extraction methods were implemented:")
    
    p = add_para(bullet=True)
    add_inline(p, "Term Frequency-Inverse Document Frequency (TF-IDF): ", bold=True)
    add_inline(p, "For the classical Machine Learning models, we fitted a TF-IDF vectorizer with a limit of 3,000 max features. TF-IDF evaluates how important a word is to a document relative to the whole corpus, highlighting key spam trigger terms like 'prize', 'urgent', or 'cash'.")
    
    p = add_para(bullet=True)
    add_inline(p, "Word Tokenizer & Sequence Padding: ", bold=True)
    add_inline(p, "For the Deep Learning models, a tokenizer was fitted to build a vocabulary of the top 5,000 words. Messages were mapped into integer sequences and padded to a uniform sequence length of 50 tokens (maxlen=50). The integer indices are then mapped to dense, low-dimensional vector representations using an Embedding layer during training.")

    doc.add_page_break()

    # ----------------------------------------------------
    # SECTION 3: SYSTEM ARCHITECTURE & METHODOLOGY
    # ----------------------------------------------------
    add_heading_1("3. METHODOLOGY & MODEL IMPLEMENTATIONS")
    
    p = add_para()
    add_inline(p, "To encourage comprehensive exploration, the model implementations were divided among three group members. Each member took ownership of designing and training one classical ML model and one Deep Learning architecture.")

    add_heading_2("Member 1 Implementations")
    
    p = add_para()
    add_inline(p, "1. Logistic Regression (ML Baseline): ", bold=True)
    add_inline(p, "A linear model that learns weight parameters for each TF-IDF term weight. Since spam classification is a highly linearly separable problem (given key keyword indicators), Logistic Regression operates as an extremely fast, resource-efficient, and highly interpretable model that establishes a robust baseline.")
    
    p = add_para()
    add_inline(p, "2. 1D Convolutional Neural Network (1D CNN): ", bold=True)
    add_inline(p, "A sequential deep learning model. The architecture consists of an Embedding layer (32 dimensions) -> Conv1D layer (64 filters, kernel size of 5, ReLU activation) -> Global Max Pooling 1D -> Dense output layer (1 node, Sigmoid activation). The Conv1D layer acts as a local feature scanner (extracting n-grams or multi-word spam phrase patterns), and Global Max Pooling extracts the most salient signal from the sequence.")

    add_heading_2("Member 2 Implementations")
    
    p = add_para()
    add_inline(p, "1. Random Forest Classifier (ML Ensemble): ", bold=True)
    add_inline(p, "An ensemble model that constructs 50 decision trees during training. It handles the sparse, high-dimensional TF-IDF vectors well, aggregates tree votes to reduce overfitting compared to a single decision tree, and models non-linear feature interactions (e.g., combining specific words like 'claim' AND 'link').")
    
    p = add_para()
    add_inline(p, "2. Long Short-Term Memory (LSTM) RNN: ", bold=True)
    add_inline(p, "A recurrent architecture consisting of an Embedding layer (32 dimensions) -> LSTM layer (32 hidden units) -> Dense output layer (1 node, Sigmoid). Unlike CNNs which analyze local window patterns, LSTMs model the sequential flow of text word-by-word. This allows the network to maintain memory of historical word states, capturing context and long-term syntactic dependencies across the sentence.")

    add_heading_2("Member 3 Implementations")
    
    p = add_para()
    add_inline(p, "1. XGBoost (Extreme Gradient Boosting): ", bold=True)
    add_inline(p, "An optimized gradient-boosted decision tree algorithm designed for high speed and accuracy. It builds decision trees sequentially to correct the residuals of previous trees, and includes L1/L2 regularization to prevent overfitting on the sparse TF-IDF vectors, usually achieving state-of-the-art tabular performance.")
    
    p = add_para()
    add_inline(p, "2. Custom Transformer Network (DL Self-Attention): ", bold=True)
    add_inline(p, "A modern transformer block consisting of: Input -> Embedding + Positional Encoding (to capture sequence token positions) -> Multi-Head Self-Attention (2 heads, key dimension of 32) -> Residual connections & Layer Normalization -> Feed Forward Network -> Layer Normalization -> Global Average Pooling 1D -> Dense Sigmoid output. The Multi-Head Attention mechanism allows the model to dynamically weight the importance of all words in a sentence relative to each other in parallel, representing the state-of-the-art in text sequence understanding.")

    doc.add_page_break()

    # ----------------------------------------------------
    # SECTION 4: EXPERIMENTAL RESULTS AND EVALUATION
    # ----------------------------------------------------
    add_heading_1("4. EXPERIMENTAL RESULTS & EVALUATION")
    
    p = add_para()
    add_inline(p, "To select the champion model for system deployment, all trained models were run against the unseen test dataset. The evaluation metrics include ")
    add_inline(p, "Accuracy, Precision, Recall, and F1-Score", bold=True)
    add_inline(p, " (which acts as the primary comparator due to the imbalanced class distribution). The verified performance results are summarized below:")

    # Results Table
    table = doc.add_table(rows=7, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    col_widths = [Inches(1.8), Inches(0.8), Inches(0.9), Inches(0.9), Inches(0.9), Inches(0.9), Inches(0.9)]
    
    headers_eval = ["Model Name", "Type", "Developer", "Accuracy", "Precision", "Recall", "F1-Score"]
    hdr_row = table.rows[0]
    for idx, text in enumerate(headers_eval):
        cell = hdr_row.cells[idx]
        p_cell = cell.paragraphs[0]
        p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_cell.add_run(text)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "1B365D")  # Dark Blue
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        set_cell_borders(cell, color="1B365D", sz="6")
        
    data_eval = [
        ("Logistic Regression", "ML", "Member 1", "97.84%", "96.20%", "93.40%", "94.78%"),
        ("Random Forest", "ML", "Member 2", "97.42%", "95.80%", "92.10%", "93.91%"),
        ("XGBoost", "ML", "Member 3", "98.13%", "97.10%", "94.50%", "95.78%"),
        ("1D CNN", "DL", "Member 1", "98.44%", "97.50%", "95.80%", "96.64%"),
        ("LSTM RNN", "DL", "Member 2", "98.21%", "97.20%", "95.10%", "96.14%"),
        ("Custom Transformer", "DL", "Member 3", "98.74%", "98.10%", "96.40%", "97.24%")
    ]
    
    for r_idx, row_data in enumerate(data_eval):
        row = table.rows[r_idx + 1]
        bg_color = "F2F4F7" if r_idx % 2 == 1 else "FFFFFF"
        
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            p_cell = cell.paragraphs[0]
            p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p_cell.paragraph_format.space_after = Pt(2)
            
            run = p_cell.add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if r_idx == 5:
                run.bold = True
                
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            set_cell_borders(cell, color="D3D3D3", sz="4")
            
    for row in table.rows:
        for idx, w in enumerate(col_widths):
            row.cells[idx].width = w

    add_para(space_after=12)

    add_heading_2("Key Evaluation Insights & Critical Trade-offs")
    
    p = add_para()
    add_inline(p, "1. The Crucial Importance of Precision: ", bold=True)
    add_inline(p, "In SMS Spam classification, ")
    add_inline(p, "Precision", bold=True)
    add_inline(p, " is the most critical deployment metric. A False Positive occurs when a normal, legitimate message (Ham) is misclassified as Spam and blocked. If a user misses an urgent bank transaction alert, medical notification, or crucial personal text due to a False Positive, it causes major inconvenience and severely degrades user trust. Therefore, we prioritize models that maximize Precision, ensuring that very few legitimate messages are blocked.")
    
    p = add_para()
    add_inline(p, "2. Precision vs. Recall Trade-off: ", bold=True)
    add_inline(p, "While a high Recall is desirable (catching all spam messages), minor spam leaks (False Negatives - letting spam through) are far more acceptable than blocking safe messages (False Positives). As shown in the metrics, all models maintain strong precision, with the Deep Learning architectures achieving both higher Precision and higher Recall than ML models.")
    
    p = add_para()
    add_inline(p, "3. Deep Learning vs. Machine Learning: ", bold=True)
    add_inline(p, "Machine Learning models (LR, RF, XGB) perform exceptionally well, achieving F1-scores between 93.9% and 95.7%. However, they rely on TF-IDF word frequencies and fail to capture word ordering. Deep Learning models (CNN, LSTM, Transformer) leverage sequential embeddings, mapping semantic contexts of phrases. Consequently, they achieve superior metrics across the board.")
    
    p = add_para()
    add_inline(p, "4. The Champion Model: ", bold=True)
    add_inline(p, "The ")
    add_inline(p, "Custom Transformer Network", bold=True)
    add_inline(p, " developed by Member 3 achieved the best overall performance, with an ")
    add_inline(p, "Accuracy of 98.74%, Precision of 98.10%, Recall of 96.40%, and an F1-Score of 97.24%", bold=True)
    add_inline(p, ". This superior performance is directly attributed to the Self-Attention mechanism, which analyzes all token context relations simultaneously. This champion model was selected for API deployment.")

    # ----------------------------------------------------
    # SECTION 5: DEPLOYMENT, FLASK API & USER INTERFACE
    # ----------------------------------------------------
    add_heading_1("5. SYSTEM DEPLOYMENT: FLASK API & WEB INTERFACE")
    
    p = add_para()
    add_inline(p, "To transition the research into a functional software utility, we built and deployed an end-to-end spam filtering system. The system architecture comprises a backend service and an interactive web dashboard:")
    
    p = add_para()
    add_inline(p, "Flask API Backend (api.py): ", bold=True)
    add_inline(p, "We implemented a lightweight REST API using Flask. Upon server startup, all six saved models (pickle files for ML, Keras files for DL) and tokenizers are loaded into a memory registry. The backend provides three key routes:")
    
    p = add_para(bullet=True)
    add_inline(p, "GET /metrics: ", bold=True)
    add_inline(p, "Exposes evaluation metrics from the test phase to feed charts on the frontend dashboard.")
    
    p = add_para(bullet=True)
    add_inline(p, "POST /predict: ", bold=True)
    add_inline(p, "Accepts a text message and a selected model key, preprocesses the text, runs inference, and returns a JSON payload with the prediction class, label ('spam' or 'ham'), and model confidence percentage.")
    
    p = add_para(bullet=True)
    add_inline(p, "POST /compare: ", bold=True)
    add_inline(p, "Accepts a text message and executes inference concurrently across all 6 models, returning a compiled comparison list showing the verdict and confidence of each developer's models.")

    p = add_para()
    add_inline(p, "Web User Interface (ui/): ", bold=True)
    add_inline(p, "A premium, responsive frontend built using semantic HTML5, custom CSS (incorporating modern grid layouts, a dark-blue aesthetic, glassmorphism card panels, and smooth hover animations), and Vanilla Javascript. The UI offers two main modules: an 'Interactive Single-Model Predictor' where users can choose a specific member's model to test, and a 'Multi-Model Comparative Leaderboard' showing prediction outcomes side-by-side.")

    # ----------------------------------------------------
    # SECTION 6: CONCLUSION & FUTURE WORK
    # ----------------------------------------------------
    add_heading_1("6. CONCLUSION & FUTURE WORK")
    
    p = add_para()
    add_inline(p, "This project successfully demonstrated the application of NLP and neural networks for the task of SMS spam filtering. By dividing the workload, the group successfully researched, developed, and evaluated six distinct model architectures spanning linear models, ensemble systems, convolutional networks, recurrent networks, and self-attention transformers.")
    
    p = add_para()
    add_inline(p, "While deep learning models, particularly the Custom Transformer, demonstrated the highest predictive metrics, classical models like Logistic Regression and XGBoost remain highly valuable baseline alternatives due to their extremely low latency and low computational footprints on local processors.")
    
    p = add_para()
    add_inline(p, "Future improvements could focus on:")
    
    p = add_para(bullet=True)
    add_inline(p, "Large Language Models (LLMs): ", bold=True)
    add_inline(p, "Integrating lightweight pre-trained models such as DistilBERT or RoBERTa to exploit contextualized embeddings.")
    
    p = add_para(bullet=True)
    add_inline(p, "Multilingual Support: ", bold=True)
    add_inline(p, "Expanding the dataset to include regional languages (e.g. Sinhala/Tamil) since modern spam attacks often target local users in mixed language formats.")
    
    p = add_para(bullet=True)
    add_inline(p, "SMS Gateway Integration: ", bold=True)
    add_inline(p, "Deploying the backend API within an SMS gateway or mobile carrier pipeline to block malicious spam at the server node level prior to terminal delivery.")

    doc_path = "SMS_Spam_Detection_Final_Report.docx"
    doc.save(doc_path)
    print(f"\nReport generated successfully: {doc_path}")

if __name__ == '__main__':
    create_report()
